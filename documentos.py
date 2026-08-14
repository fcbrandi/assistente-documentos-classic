from datetime import datetime
from pathlib import Path
import json

from docx import Document
from openai import OpenAI
from pypdf import PdfReader

from config import (
    ARQUIVO_ESTRUTURAS,
    ARQUIVO_STATUS,
    FORMATOS_ACEITOS,
    MODELO_ANALISE_ESTRUTURA,
    PASTA_UPLOADS,
)


def carregar_desativados():
    if not ARQUIVO_STATUS.exists():
        return set()

    try:
        return set(json.loads(ARQUIVO_STATUS.read_text(encoding="utf-8")))
    except Exception:
        return set()


def salvar_desativados(desativados):
    ARQUIVO_STATUS.write_text(
        json.dumps(sorted(desativados), ensure_ascii=False),
        encoding="utf-8",
    )


def listar_documentos():
    desativados = carregar_desativados()
    documentos = []

    for arquivo in PASTA_UPLOADS.iterdir():
        if arquivo.is_file() and arquivo.suffix.lower() in FORMATOS_ACEITOS:
            documentos.append(
                {
                    "nome": arquivo.name,
                    "tipo": arquivo.suffix.upper().replace(".", ""),
                    "data": datetime.fromtimestamp(
                        arquivo.stat().st_mtime
                    ).strftime("%d/%m/%Y %H:%M"),
                    "ativo": arquivo.name not in desativados,
                }
            )

    return sorted(
        documentos,
        key=lambda documento: documento["data"],
        reverse=True,
    )


def documentos_ativos():
    desativados = carregar_desativados()

    return sorted(
        [
            arquivo
            for arquivo in PASTA_UPLOADS.iterdir()
            if (
                arquivo.is_file()
                and arquivo.suffix.lower() in FORMATOS_ACEITOS
                and arquivo.name not in desativados
            )
        ],
        key=lambda arquivo: arquivo.stat().st_mtime,
    )


def extrair_texto(caminho):
    extensao = caminho.suffix.lower()

    if extensao == ".txt":
        return caminho.read_text(encoding="utf-8", errors="replace")

    if extensao == ".docx":
        documento = Document(str(caminho))
        partes = [paragrafo.text for paragrafo in documento.paragraphs]

        for tabela in documento.tables:
            for linha in tabela.rows:
                partes.append(
                    " | ".join(celula.text for celula in linha.cells)
                )

        return "\n".join(partes)

    if extensao == ".pdf":
        leitor = PdfReader(str(caminho))
        return "\n".join(
            pagina.extract_text() or ""
            for pagina in leitor.pages
        )

    raise ValueError("Formato de arquivo não aceito.")


def carregar_estruturas():
    if not ARQUIVO_ESTRUTURAS.exists():
        return {}

    try:
        return json.loads(
            ARQUIVO_ESTRUTURAS.read_text(encoding="utf-8")
        )
    except Exception:
        return {}


def salvar_estruturas(estruturas):
    ARQUIVO_ESTRUTURAS.write_text(
        json.dumps(estruturas, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def extrair_json(texto):
    inicio = texto.find("{")
    fim = texto.rfind("}")

    if inicio < 0 or fim < 0:
        raise ValueError("A análise não retornou um JSON válido.")

    return json.loads(texto[inicio : fim + 1])


def resumo_paginas(leitor):
    total_paginas = len(leitor.pages)

    if total_paginas <= 30:
        numeros = list(range(1, total_paginas + 1))
    else:
        numeros = list(range(1, 16))
        numeros += list(range(total_paginas - 14, total_paginas + 1))

    partes = []

    for numero_pagina in numeros:
        texto = leitor.pages[numero_pagina - 1].extract_text() or ""
        texto = " ".join(texto.split())

        partes.append(
            f"PÁGINA {numero_pagina}:\n{texto[:700]}"
        )

    return "\n\n".join(partes)


def analisar_estrutura_pdf(caminho, leitor):
    total_paginas = len(leitor.pages)
    resumo = resumo_paginas(leitor)

    prompt = f"""
Você analisa a estrutura de um PDF antes de ele entrar em uma biblioteca.

Identifique apenas o corpo principal do documento.
Desconsidere capa, notas introdutórias, prefácios, comentários editoriais,
sumários, questionários, exercícios, índices e anexos.

Com base nas páginas resumidas abaixo, responda SOMENTE em JSON válido,
neste formato:

{{
  "pagina_inicial": 1,
  "pagina_final": {total_paginas},
  "marcador_inicio": "",
  "observacao": ""
}}

Regras:
- "pagina_inicial" é a primeira página que contém o texto principal.
- "pagina_final" é a última página do texto principal.
- Use "marcador_inicio" apenas se, na primeira página útil, houver texto
  editorial antes do título ou do início principal. Nesse caso, informe um
  título ou frase exata que marque o começo do conteúdo principal.
- Se não houver certeza, mantenha todas as páginas e deixe
  "marcador_inicio" vazio.

ARQUIVO:
{caminho.name}

PÁGINAS:
{resumo}
"""

    cliente = OpenAI()
    resultado = cliente.responses.create(
        model=MODELO_ANALISE_ESTRUTURA,
        input=prompt,
    )

    dados = extrair_json(resultado.output_text)

    try:
        pagina_inicial = int(dados.get("pagina_inicial", 1))
        pagina_final = int(dados.get("pagina_final", total_paginas))
    except (TypeError, ValueError):
        pagina_inicial = 1
        pagina_final = total_paginas

    pagina_inicial = max(1, min(pagina_inicial, total_paginas))
    pagina_final = max(pagina_inicial, min(pagina_final, total_paginas))

    return {
        "pagina_inicial": pagina_inicial,
        "pagina_final": pagina_final,
        "marcador_inicio": str(
            dados.get("marcador_inicio", "")
        ).strip(),
        "observacao": str(dados.get("observacao", "")).strip(),
    }


def obter_estrutura_pdf(caminho, leitor):
    estruturas = carregar_estruturas()
    modificado_em = caminho.stat().st_mtime_ns
    registro = estruturas.get(caminho.name)

    if registro and registro.get("modificado_em") == modificado_em:
        return registro

    estrutura = analisar_estrutura_pdf(caminho, leitor)
    estrutura["modificado_em"] = modificado_em

    estruturas[caminho.name] = estrutura
    salvar_estruturas(estruturas)

    return estrutura


def extrair_partes_referenciadas(caminho):
    extensao = caminho.suffix.lower()

    if extensao == ".pdf":
        leitor = PdfReader(str(caminho))
        estrutura = obter_estrutura_pdf(caminho, leitor)

        pagina_inicial = estrutura["pagina_inicial"]
        pagina_final = estrutura["pagina_final"]
        marcador_inicio = estrutura["marcador_inicio"]
        partes = []

        for numero_pagina, pagina in enumerate(leitor.pages, start=1):
            if numero_pagina < pagina_inicial:
                continue

            if numero_pagina > pagina_final:
                continue

            texto = pagina.extract_text() or ""

            if numero_pagina == pagina_inicial and marcador_inicio:
                posicao = texto.lower().find(marcador_inicio.lower())

                if posicao >= 0:
                    texto = texto[
                        posicao + len(marcador_inicio):
                    ]

            if texto.strip():
                partes.append(
                    {
                        "pagina": numero_pagina,
                        "texto": texto,
                    }
                )

        return partes

    return [
        {
            "pagina": None,
            "texto": extrair_texto(caminho),
        }
    ]